"""Tests for ToolRegistry and ReadFileTool."""

import asyncio
from pathlib import Path

import pytest

from mybot.agent.tools.filesystem import ReadFileTool
from mybot.agent.tools.registry import ToolRegistry


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def run(coro):
    """Run an async coroutine from sync test code."""
    return asyncio.get_event_loop().run_until_complete(coro)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture()
def registry() -> ToolRegistry:
    reg = ToolRegistry()
    reg.register(ReadFileTool())
    return reg


@pytest.fixture()
def sample_txt(tmp_path: Path) -> Path:
    p = tmp_path / "hello.txt"
    p.write_text("line1\nline2\nline3\n", encoding="utf-8")
    return p


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal valid .docx file."""
    from docx import Document

    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello DOCX")
    doc.add_paragraph("Second paragraph")
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# ToolRegistry: register / unregister / get / has
# ---------------------------------------------------------------------------

class TestToolRegistryBasic:
    def test_register_and_get(self):
        reg = ToolRegistry()
        tool = ReadFileTool()
        reg.register(tool)
        assert reg.get("read_file") is tool

    def test_has(self, registry: ToolRegistry):
        assert registry.has("read_file") is True
        assert registry.has("nonexistent") is False

    def test_unregister(self, registry: ToolRegistry):
        registry.unregister("read_file")
        assert registry.has("read_file") is False
        assert registry.get("read_file") is None

    def test_unregister_nonexistent_is_noop(self):
        reg = ToolRegistry()
        reg.unregister("fake")  # should not raise


# ---------------------------------------------------------------------------
# ToolRegistry: get_definition / get_tool_schemas
# ---------------------------------------------------------------------------

class TestToolDefinition:
    def test_get_definition_returns_list(self, registry: ToolRegistry):
        defs = registry.get_definition()
        assert isinstance(defs, list)
        assert len(defs) == 1

    def test_definition_has_openai_format(self, registry: ToolRegistry):
        schema = registry.get_definition()[0]
        assert schema["type"] == "function"
        assert "function" in schema
        fn = schema["function"]
        assert fn["name"] == "read_file"
        assert "description" in fn
        assert "parameters" in fn

    def test_get_tool_schemas_alias(self, registry: ToolRegistry):
        assert registry.get_tool_schemas() == registry.get_definition()

    def test_cache_invalidation(self, registry: ToolRegistry):
        defs1 = registry.get_definition()
        registry.register(ReadFileTool())  # re-register triggers cache clear
        defs2 = registry.get_definition()
        assert len(defs2) == 1  # same tool, but cache was rebuilt


# ---------------------------------------------------------------------------
# ToolRegistry: prepare_call
# ---------------------------------------------------------------------------

class TestPrepareCall:
    def test_valid_call(self, registry: ToolRegistry):
        tool, params, err = registry.prepare_call("read_file", {"path": "/tmp/x"})
        assert tool is not None
        assert params == {"path": "/tmp/x"}
        assert err is None

    def test_missing_tool(self, registry: ToolRegistry):
        tool, params, err = registry.prepare_call("fake_tool", {"a": 1})
        assert tool is None
        assert err is not None
        assert "not found" in err

    def test_non_dict_params(self, registry: ToolRegistry):
        tool, params, err = registry.prepare_call("read_file", "bad")
        assert err is not None
        assert "JSON object" in err

    def test_coerce_json_string_params(self, registry: ToolRegistry):
        tool, params, err = registry.prepare_call("read_file", '{"path": "/tmp/x"}')
        assert err is None
        assert params == {"path": "/tmp/x"}


# ---------------------------------------------------------------------------
# ToolRegistry: execute
# ---------------------------------------------------------------------------

class TestExecute:
    def test_execute_missing_file(self, registry: ToolRegistry):
        result = run(registry.execute("read_file", {"path": "/tmp/__no_such_file__.txt"}))
        assert "Error" in result or "not found" in result.lower()

    def test_execute_missing_tool(self, registry: ToolRegistry):
        result = run(registry.execute("fake_tool", {}))
        assert "Error" in result
        assert "not found" in result

    def test_execute_reads_file(self, registry: ToolRegistry, sample_txt: Path):
        result = run(registry.execute("read_file", {"path": str(sample_txt)}))
        assert "line1" in result
        assert "line2" in result


# ---------------------------------------------------------------------------
# ReadFileTool: direct tests
# ---------------------------------------------------------------------------

class TestReadFileTool:
    def test_read_text_file(self, sample_txt: Path):
        tool = ReadFileTool()
        result = run(tool.execute(path=str(sample_txt)))
        assert "line1" in result
        assert "line3" in result

    def test_read_nonexistent(self):
        tool = ReadFileTool()
        result = run(tool.execute(path="/tmp/__no_such_file__.txt"))
        assert "Error" in result or "not found" in result.lower()

    def test_read_no_path(self):
        tool = ReadFileTool()
        result = run(tool.execute(path=None))
        assert "Error" in result

    def test_read_empty_path(self):
        tool = ReadFileTool()
        result = run(tool.execute(path=""))
        assert "Error" in result

    def test_read_docx(self, sample_docx: Path):
        tool = ReadFileTool()
        result = run(tool.execute(path=str(sample_docx)))
        assert "Hello DOCX" in result
        assert "Second paragraph" in result

    def test_read_with_offset_and_limit(self, sample_txt: Path):
        tool = ReadFileTool()
        result = run(tool.execute(path=str(sample_txt), offset=2, limit=1))
        assert "line2" in result
        assert "line1" not in result

    def test_tool_properties(self):
        tool = ReadFileTool()
        assert tool.name == "read_file"
        assert "description" in dir(tool)
        assert tool.read_only is True

    def test_parameters_schema(self):
        tool = ReadFileTool()
        params = tool.parameters
        assert params["type"] == "object"
        assert "path" in params["properties"]
        assert "path" in params["required"]


# ---------------------------------------------------------------------------
# Tool: to_schema / cast_params
# ---------------------------------------------------------------------------

class TestToolMethods:
    def test_to_schema(self):
        tool = ReadFileTool()
        schema = tool.to_schema()
        assert schema["type"] == "function"
        assert schema["function"]["name"] == "read_file"

    def test_cast_params_identity(self):
        tool = ReadFileTool()
        params = {"path": "/tmp/test.txt", "offset": 1}
        result = tool.cast_params(params)
        assert result["path"] == "/tmp/test.txt"
        assert result["offset"] == 1

    def test_cast_params_coerces_string_int(self):
        tool = ReadFileTool()
        params = {"path": "/tmp/test.txt", "offset": "5"}
        result = tool.cast_params(params)
        assert result["offset"] == 5
        assert isinstance(result["offset"], int)
